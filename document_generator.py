"""
Document Generator Module
Generates .docx files with proper formatting, images, and citations
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

class DocumentGenerator:
    def __init__(self, settings):
        self.settings = settings
        self.citation_style = settings.get('citation_style', 'APA')
        self.include_images = settings.get('include_images', True)
        self.output_format = settings.get('output_format', 'docx')
        
    def generate_docx(self, content, topic, images=None):
        """
        Generate a formatted .docx document
        
        Args:
            content: Dictionary containing document sections and text
            topic: Main topic title
            images: List of image objects to include
        
        Returns:
            Path to generated .docx file
        """
        doc = Document()
        
        # Set document styles
        self._setup_styles(doc)
        
        # Add title page
        self._add_title_page(doc, topic)
        
        # Add table of contents (placeholder)
        self._add_table_of_contents(doc)
        
        # Add content sections
        if isinstance(content, dict) and 'sections' in content:
            for i, section in enumerate(content['sections']):
                self._add_section(doc, section, i)
                
                # Add images after each section
                if images and self.include_images:
                    section_images = [img for img in images 
                                    if img.get('section_id', '').startswith(section.get('id', ''))]
                    self._add_images_to_section(doc, section_images)
        
        elif isinstance(content, str):
            # If content is a string, add it directly
            self._add_text_content(doc, content)
            
            # Add images throughout
            if images and self.include_images:
                self._distribute_images(doc, images)
        
        # Add references/bibliography
        if 'references' in content:
            self._add_references(doc, content['references'])
        
        # Add appendix if exists
        if 'appendix' in content:
            self._add_appendix(doc, content['appendix'])
        
        # Save document
        filename = self._generate_filename(topic)
        filepath = os.path.join('uploads', 'generated', filename)
        
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        doc.save(filepath)
        
        return filepath
    
    def _setup_styles(self, doc):
        """Set up document styles"""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Heading styles
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.color.rgb = RGBColor(0, 51, 102)
    
    def _add_title_page(self, doc, topic):
        """Add a professional title page"""
        # Add spacing
        for _ in range(4):
            doc.add_paragraph()
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(topic)
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Generated Research Document')
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(102, 102, 102)
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
        run.font.size = Pt(11)
        
        # Add page break
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc):
        """Add table of contents (placeholder)"""
        toc_heading = doc.add_heading('Table of Contents', level=1)
        doc.add_paragraph('[Table of Contents will be generated here]')
        doc.add_page_break()
    
    def _add_section(self, doc, section, index):
        """Add a content section with proper formatting"""
        # Section heading
        heading_text = section.get('title', f'Section {index + 1}')
        doc.add_heading(heading_text, level=1)
        
        # Section content
        content_text = section.get('content', '')
        
        # Split content into paragraphs
        paragraphs = content_text.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.style = doc.styles['Normal']
        
        # Add spacing after section
        doc.add_paragraph()
    
    def _add_images_to_section(self, doc, images):
        """Add images to the document section"""
        if not images:
            return
        
        for img_data in images:
            try:
                # Image path
                img_path = img_data.get('url', '/static/images/placeholder.jpg')
                
                # Check if file exists locally
                if os.path.exists(img_path.lstrip('/')):
                    # Add image with caption
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    run = paragraph.add_run()
                    run.add_picture(img_path.lstrip('/'), width=Inches(5.5))
                    
                    # Add caption
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption.add_run(img_data.get('caption', 'Image'))
                    caption_run.italic = True
                    caption_run.font.size = Pt(9)
                    
                    # Add spacing
                    doc.add_paragraph()
                    
            except Exception as e:
                print(f"Error adding image: {e}")
                # Add placeholder text if image can't be added
                doc.add_paragraph(f"[Image: {img_data.get('caption', 'Image placeholder')}]")
    
    def _add_text_content(self, doc, content):
        """Add plain text content"""
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
    
    def _distribute_images(self, doc, images):
        """Distribute images throughout the document"""
        if not images:
            return
        
        # Get total paragraphs
        paragraphs = doc.paragraphs
        total_paragraphs = len(paragraphs)
        
        if total_paragraphs > 0 and images:
            # Calculate spacing for images
            spacing = max(1, total_paragraphs // len(images))
            
            for i, img_data in enumerate(images):
                try:
                    # Insert at calculated positions
                    position = (i + 1) * spacing
                    if position < total_paragraphs:
                        paragraph = paragraphs[position]
                        
                        # Insert image before this paragraph
                        img_paragraph = doc.add_paragraph()
                        img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add placeholder if image file doesn't exist
                        img_paragraph.add_run(f"[Image {i+1}: {img_data.get('caption', '')}]")
                        
                except Exception as e:
                    print(f"Error distributing image: {e}")
    
    def _add_references(self, doc, references):
        """Add references/bibliography section"""
        doc.add_page_break()
        doc.add_heading('References', level=1)
        
        if isinstance(references, list):
            for ref in references:
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                
                if self.citation_style == 'APA':
                    p.add_run(ref.get('apa', str(ref))).font.size = Pt(10)
                elif self.citation_style == 'MLA':
                    p.add_run(ref.get('mla', str(ref))).font.size = Pt(10)
                else:
                    p.add_run(str(ref)).font.size = Pt(10)
    
    def _add_appendix(self, doc, appendix_content):
        """Add appendix section"""
        doc.add_page_break()
        doc.add_heading('Appendix', level=1)
        
        if isinstance(appendix_content, str):
            doc.add_paragraph(appendix_content)
        elif isinstance(appendix_content, list):
            for item in appendix_content:
                doc.add_paragraph(str(item))
    
    def _generate_filename(self, topic):
        """Generate a clean filename"""
        # Clean topic for filename
        clean_topic = "".join(c for c in topic[:50] if c.isalnum() or c.isspace())
        clean_topic = clean_topic.replace(' ', '_').lower()
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        return f"{clean_topic}_{timestamp}.docx"
